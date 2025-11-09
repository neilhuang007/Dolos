"""Document builder for creating DOCX files with track changes."""

from datetime import datetime
from pathlib import Path
from typing import List, Optional

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import Document, Sentence


class DocumentBuilder:
    """Build Word documents with metadata and track changes."""

    def __init__(self):
        """Initialize document builder."""
        self.document = None

    def create_document(
        self,
        sentences: List[Sentence],
        output_path: str,
        author: str = "Dolos",
        title: Optional[str] = None,
        subject: Optional[str] = None,
        keywords: Optional[str] = None,
        comments: Optional[str] = None,
        total_edit_minutes: Optional[int] = None
    ) -> str:
        """Create a DOCX document from sentences.

        Args:
            sentences: List of Sentence objects with metadata
            output_path: Path to save the document
            author: Document author
            title: Document title
            subject: Document subject
            keywords: Document keywords/tags
            comments: Document comments
            total_edit_minutes: Total editing time in minutes

        Returns:
            Path to created document
        """
        # Create new document
        self.document = DocxDocument()

        # Set core properties
        core_props = self.document.core_properties
        core_props.author = author
        core_props.last_modified_by = author

        if title:
            core_props.title = title

        if subject:
            core_props.subject = subject

        if keywords:
            core_props.keywords = keywords

        if comments:
            core_props.comments = comments

        # Set timestamps from first and last sentence
        if sentences:
            core_props.created = sentences[0].created_timestamp
            core_props.modified = sentences[-1].modified_timestamp

        # Note: Total editing time will be set via XML manipulation after document is saved
        # python-docx doesn't expose this property directly

        # Add content (we'll inject track changes later via XML)
        # For now, just add the sentences as regular paragraphs
        for sentence in sentences:
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run(sentence.sentence_text)

            # Add space after each sentence except the last
            if sentence != sentences[-1]:
                paragraph.add_run(" ")

        # Save the document
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(str(output_path))

        return str(output_path)

    def create_simple_document(
        self,
        text: str,
        output_path: str,
        author: str = "Dolos",
        created_time: Optional[datetime] = None,
        modified_time: Optional[datetime] = None
    ) -> str:
        """Create a simple DOCX document without track changes.

        Args:
            text: Document text content
            output_path: Path to save the document
            author: Document author
            created_time: Creation timestamp
            modified_time: Modification timestamp

        Returns:
            Path to created document
        """
        self.document = DocxDocument()

        # Set core properties
        core_props = self.document.core_properties
        core_props.author = author
        core_props.last_modified_by = author

        if created_time:
            core_props.created = created_time
        if modified_time:
            core_props.modified = modified_time

        # Add content
        paragraph = self.document.add_paragraph(text)

        # Save
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(str(output_path))

        return str(output_path)
